# src/export_to_word.py
import sqlite3
import pandas as pd
from docx import Document
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows

def export_to_word(db_name, table_name, output_filename, selected_columns=None):
    conn = sqlite3.connect(db_name)
    if selected_columns:
        columns = ', '.join(selected_columns)
        query = f"SELECT {columns} FROM {table_name}"
    else:
        query = f"SELECT * FROM {table_name}"
    df = pd.read_sql_query(query, conn)
    doc = Document()
    doc.add_heading(table_name, level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    for i, col_name in enumerate(df.columns):
        table.cell(0, i).text = col_name
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row):
            row_cells[i].text = str(cell_value)
    doc.save(output_filename)
    conn.close()


def export_to_excel(db_name, table_name, output_filename, selected_columns=None):
    conn = sqlite3.connect(db_name)
    if selected_columns:
        columns = ', '.join(selected_columns)
        query = f"SELECT {columns} FROM {table_name}"
    else:
        query = f"SELECT * FROM {table_name}"
    df = pd.read_sql_query(query, conn)
    
    wb = Workbook()
    ws = wb.active
    
    # Write the header
    header = df.columns.tolist()
    ws.append(header)
    
    # Write data rows
    for row in dataframe_to_rows(df, index=False, header=False):
        ws.append(row)
    
    wb.save(output_filename)
    conn.close()
    
    
def export_to_csv(db_name, table_name, output_filename, selected_columns=None):
    conn = sqlite3.connect(db_name)
    if selected_columns:
        columns = ', '.join(selected_columns)
        query = f"SELECT {columns} FROM {table_name}"
    else:
        query = f"SELECT * FROM {table_name}"
    df = pd.read_sql_query(query, conn)
    df.to_csv(output_filename, index=False)
    conn.close()


def export_to_json(db_name, table_name, output_filename, selected_columns=None):
    conn = sqlite3.connect(db_name)
    if selected_columns:
        columns = ', '.join(selected_columns)
        query = f"SELECT {columns} FROM {table_name}"
    else:
        query = f"SELECT * FROM {table_name}"
    df = pd.read_sql_query(query, conn)
    json_data = df.to_json(orient="records", indent=4)
    with open(output_filename, 'w') as json_file:
        json_file.write(json_data)
    conn.close()